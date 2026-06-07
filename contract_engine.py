import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import datetime

def generar_contrato_opcion_compra(designacion_catastral, propietario, comprador, analisis_texto):
    """
    Genera un documento Word (.docx) formal con un contrato de Promesa de Venta 
    blindado bajo la Ley 108-05 y con cláusulas de mitigación dinámicas.
    """
    doc = docx.Document()
    
    # Configuración de márgenes estándar legal
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Estilo de fuente global (Times New Roman o Arial son el estándar legal en RD)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # TÍTULO PRINCIPAL
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = titulo.add_run("ACTO BAJO FIRMA PRIVADA\nCONTRATO DE PROMESA DE VENTA Y OPCIÓN A COMPRA")
    run_t.bold = True
    run_t.size = Pt(14)
    doc.add_paragraph("\n")

    # PREÁMBULO DE LAS PARTES
    p_partes = doc.add_paragraph()
    p_partes.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_partes.add_run("REUNIDOS de una parte, el señor/sociedad ").font.name = 'Times New Roman'
    p_partes.add_run(f"{propietario}").bold = True
    p_partes.add_run(", en lo adelante denominado como 'EL PROMETIENTE VENDEDOR'; y de la otra parte, el señor/sociedad ")
    p_partes.add_run(f"{comprador}").bold = True
    p_partes.add_run(", en lo adelante denominado como 'EL PROMETIENTE COMPRADOR'. Se ha convenido y pactado el siguiente contrato:")
    
    doc.add_paragraph("\n")

    # CLÁUSULA PRIMERA: OBJETO E INMUEBLE
    doc.add_paragraph("CLÁUSULA PRIMERA: OBJETO.").bold = True
    p_objeto = doc.add_paragraph()
    p_objeto.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_objeto.add_run("EL PROMETIENTE VENDEDOR se compromete a vender, y EL PROMETIENTE COMPRADOR se compromete a comprar el inmueble debidamente amparado por la normativa de la Ley No. 108-05 de Registro Inmobiliario, identificado como: ")
    p_objeto.add_run(f"{designacion_catastral}.").bold = True

    # CLÁUSULA SEGUNDA: CONDICIÓN DINÁMICA DE MITIGACIÓN (El cerebro del sistema)
    doc.add_paragraph("\nCLÁUSULA SEGUNDA: CONDICIONES ESPECIALES DE MITIGACIÓN DE RIESGOS.").bold = True
    p_mitigacion = doc.add_paragraph()
    p_mitigacion.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Evaluamos qué riesgo encontró la IA para inyectar la cláusula adecuada
    if "Constancia Anotada" in analisis_texto or "790-2022" in analisis_texto:
        p_mitigacion.add_run(
            "PARÁGRAFO ESPECIAL (OBLIGATORIEDAD DE DESLINDE - RESOLUCIÓN 790-2022): Las partes reconocen expresamente "
            "que el inmueble objeto del presente contrato se encuentra sustentado sobre una Constancia Anotada (Carta de Constancia). "
            "Por tanto, es condición resolutoria y obligatoria para la firma del Acto de Venta Definitivo que EL PROMETIENTE VENDEDOR "
            "culmine a sus propios gastos el proceso de DESLINDE de conformidad con la Resolución No. 790-2022 ante la Dirección Regional "
            "de Mensuras Catastrales. Se establece la retención del sesenta por ciento (60%) del precio de venta pactado, el cual será "
            "entregado únicamente tras la emisión del Certificado de Título debidamente individualizado por el Registrador de Títulos correspondiente."
        )
    elif "Riesgo Crítico" in analisis_texto or "Oposición" in analisis_texto or "Litis" in analisis_texto:
        p_mitigacion.add_run(
            "PARÁGRAFO ESPECIAL (BLOQUEO Y SANEAMIENTO DE LITIS / CARGAS): Las partes hacen constar que sobre el inmueble "
            "pesa una anotación o contingencia legal identificada en la auditoría algorítmica de riesgos. EL PROMETIENTE VENDEDOR "
            "se obliga formalmente a interponer todas las acciones de saneamiento y levantamiento de cargas de conformidad con el "
            "Reglamento General de Tribunales de Tierras (Res. 787-2022), fijando un plazo de noventa (90) días para la entrega del "
            "inmueble completamente libre de litigios."
        )
    else:
        p_mitigacion.add_run(
            "PARÁGRAFO ESPECIAL (SITUACIÓN ESTÁNDAR): EL PROMETIENTE VENDEDOR garantiza que el inmueble se encuentra "
            "debidamente deslindado conforme a la Resolución No. 789-2022 y libre de toda carga, gravamen o litis sobre derechos "
            "registrados, comprometiéndose a mantener dicha condición pacífica hasta la firma definitiva."
        )

    # CLÁUSULA TERCERA: JURISDICCIÓN COMPETENTE
    doc.add_paragraph("\nCLÁUSULA TERCERA: JURISDICCIÓN COMPETENTE.").bold = True
    p_juris = doc.add_paragraph()
    p_juris.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_juris.add_run(
        "Para todo lo no previsto en el presente contrato, las partes se remiten de forma principal a las disposiciones de la "
        "Ley No. 108-05 de Registro Inmobiliario y, de forma supletoria, al Derecho Común. Cualquier diferendo será llevado ante el "
        "Tribunal de Tierras de Jurisdicción Original territorialmente competente."
    )

    doc.add_paragraph("\n\nHecho en dos (02) originales de un mismo tenor y efecto, en la Ciudad de Santo Domingo, República Dominicana, a los " + datetime.date.today().strftime('%d días del mes %m del año %Y') + ".")

    # Guardar el documento en memoria buffer binaria para que Streamlit pueda descargarlo
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()